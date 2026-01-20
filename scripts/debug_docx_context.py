
import docx
import os
import re

docx_path = '102 พระธัมมปทัฏฐกถา ภาค ๒ (บาลี-โดยอรรถ).docx'
doc = docx.Document(docx_path)

# Previous sentence ending
target = "น ทิฏฺฐปุพฺโพติ อาห"
print(f"Searching for end of previous sentence: '{target}'...")

found_index = -1
for i, para in enumerate(doc.paragraphs):
    if target in para.text:
        print(f"Found match at paragraph {i}:")
        print(f"Text: {para.text}")
        found_index = i
        break

if found_index != -1:
    print("\n--- Next 5 Paragraphs ---")
    for j in range(1, 6):
        if found_index + j < len(doc.paragraphs):
            print(f"[{found_index + j}]: {doc.paragraphs[found_index + j].text}")
else:
    print("Previous sentence target not found.")
