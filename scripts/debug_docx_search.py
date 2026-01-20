
import docx
import os
import sys

docx_path = '102 พระธัมมปทัฏฐกถา ภาค ๒ (บาลี-โดยอรรถ).docx'
doc = docx.Document(docx_path)

target = sys.argv[1] if len(sys.argv) > 1 else "เอวํ ชีวนฺตสฺส"
print(f"Searching for '{target}'...")

found = False
for i, para in enumerate(doc.paragraphs):
    if target in para.text:
        print(f"Found at paragraph {i}:")
        print(f"Text: {para.text}")
        # Print surrounding paragraphs for context
        start = max(0, i - 2)
        end = min(len(doc.paragraphs), i + 3)
        for j in range(start, end):
             print(f"  Para {j}: {doc.paragraphs[j].text}")
        found = True
        
if not found:
    print("Not found.")
