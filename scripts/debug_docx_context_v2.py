
import docx
import os

docx_path = '102 พระธัมมปทัฏฐกถา ภาค ๒ (บาลี-โดยอรรถ).docx'
doc = docx.Document(docx_path)

# Try a very unique word
target = "จตฺตาฬีสโกฏิวิภโว"
print(f"Searching for unique word: '{target}'...")

found_index = -1
for i, para in enumerate(doc.paragraphs):
    if target in para.text:
        print(f"Found match at paragraph {i}:")
        print(f"Text: {para.text}")
        found_index = i
        break

if found_index != -1:
    print("\n--- Next 10 Paragraphs ---")
    for j in range(1, 11):
        if found_index + j < len(doc.paragraphs):
            print(f"[{found_index + j}]: {doc.paragraphs[found_index + j].text}")
else:
    print("Target not found.")
